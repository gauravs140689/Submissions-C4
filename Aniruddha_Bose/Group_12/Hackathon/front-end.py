"""
front-end.py
============
Smart Medical Research Assistant — Gradio UI

A production-grade interface for medical research, featuring 
multi-agent pipeline execution and RAG-based knowledge retrieval.

Requirements:
    gradio>=4.0, python-env, python-docx, lancedb, langchain-openai, langchain-community
"""

import os
import sys
import io
import re
import time
import threading
from datetime import datetime
from pathlib import Path

import gradio as gr
from gradio_rich_textbox import RichTextbox
from dotenv import load_dotenv
from docx import Document
import lancedb
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import LanceDB

# ── Load environment variables ────────────────────────────────────────────────
load_dotenv()

# ── Import multi-agent.py — DO NOT MODIFY ORIGINAL ────────────────────────────
import importlib.util

def import_module_from_path(module_name, file_path):
    spec = importlib.util.spec_from_file_location(module_name, file_path)
    if spec is None or spec.loader is None:
        raise FileNotFoundError(f"Missing {file_path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module

try:
    multi_agent = import_module_from_path("multi_agent", "multi-agent.py")
    agent_app = multi_agent.app
    AgentState = multi_agent.AgentState
except Exception as e:
    print(f"Error loading multi-agent.py: {e}")
    sys.exit(1)

# ── Import RAG engine — DO NOT MODIFY ORIGINAL ────────────────────────────────
import rag

# ── Setup ───────────────────────────────────────────────────────────────────
REPORTS_DIR = Path("reports")
REPORTS_DIR.mkdir(exist_ok=True)

# ── RAG Helpers ─────────────────────────────────────────────────────────────
def chunk_and_vectorize(directory):
    docs = rag.load_documents(str(directory))
    if not docs: return False
    rag.setup_vector_db(docs)
    return True

def query_vector_db(query):
    try:
        embeddings = OpenAIEmbeddings(
            model=rag.EMBEDDING_MODEL_NAME,
            openai_api_key=rag.OPENROUTER_API_KEY,
            openai_api_base=rag.OPENROUTER_BASE_URL,
            check_embedding_ctx_length=False
        )
        db = lancedb.connect(rag.VECTOR_DB_NAME)
        table_name = "vectors"
        vectorstore = LanceDB(connection=db, embedding=embeddings, table_name=table_name)
        chain = rag.get_rag_chain(vectorstore)
        return chain.invoke(query)
    except Exception as e:
        return f"Error: {e}"

# ── CSS ─────────────────────────────────────────────────────────────────────
# We specify heights in vh to ensure it fits on screen without scrolling.
CUSTOM_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

:root {
  --bg-color: #ffffff;
  --title-color: #1f2937;
  --text-color: #2563eb; /* Primary text: Attractive blue */
  --border-color: #e5e7eb;
}

.dark {
  --bg-color: #111827;
  --title-color: #ffffff; /* Title becomes white */
  --text-color: #93c5fd; /* Vibrant light blue/color for other text */
  --border-color: #374151;
}

body, .gradio-container {
  background: var(--bg-color) !important;
  font-family: 'Inter', sans-serif !important;
  margin: 0 !important;
  padding: 0 !important;
  height: 100vh;
  overflow: hidden;
}

#header-row {
  height: 8vh;
  display: flex !important;
  align-items: center;
  justify-content: center;
  border-bottom: 2px solid var(--border-color);
  position: relative;
}

#title-markdown {
  margin: 0 !important;
  padding: 0 !important;
}

#title-markdown h1 {
  color: var(--title-color) !important;
  font-size: 28px !important;
  font-weight: 700 !important;
  margin: 0 !important;
  text-align: center;
}

.theme-toggle-btn {
  position: absolute;
  right: 20px;
  top: 50%;
  transform: translateY(-50%);
  background: transparent !important;
  border: none !important;
  font-size: 24px !important;
  cursor: pointer !important;
  padding: 0 !important;
}

.main-content-row { height: 60vh; gap: 20px; padding: 10px 20px; }
.inputs-row { height: 15vh; gap: 20px; padding: 10px 20px; }
.buttons-row { height: 10vh; gap: 20px; padding: 10px 20px; }

/* Textboxes styling */
textarea {
  color: var(--text-color) !important;
  background: var(--bg-color) !important;
  border: 1px solid var(--border-color) !important;
  font-size: 14px !important;
  line-height: 1.5 !important;
  border-radius: 8px !important;
}

/* Specific heights and scrollbars */
#t1 textarea { height: 40vh !important; }
#t2 textarea { height: 10vh !important; }
#t3 textarea, #t5 textarea { height: 10vh !important; }
#t4 textarea { height: 56vh !important; }

/* Handle placeholder color */
textarea::placeholder { color: #9ca3af !important; }

/* Button styling */
.primary-btn { background: #2563eb !important; color: white !important; }
.secondary-btn { background: #10b981 !important; color: white !important; }
.secondary-btn:disabled { background: #9ca3af !important; color: #f3f4f6 !important; cursor: not-allowed; }
.arrow-btn { font-size: 16px !important; }

/* Hide labels to keep it clean if desired, or color them */
label span { color: var(--text-color) !important; font-weight: 600 !important; text-transform: uppercase; font-size: 15px; }
"""

# ── Logic ───────────────────────────────────────────────────────────────────

def run_pipeline(query):
    if not query:
        yield "", "", gr.update(interactive=False)
        return
    
    # Start sequence
    initial_state = AgentState(
        user_query=query,
        is_valid_query=False,
        context_data={},
        retry_count=0,
        critic_summary={},
        insight_summary="",
        final_report={}
    )
    
    t1_acc = ""
    t2_acc = ""
    
    # Capture stdout for logs (T2)
    log_stream = io.StringIO()
    old_stdout = sys.stdout
    sys.stdout = log_stream
    
    try:
        for update in agent_app.stream(initial_state, stream_mode="updates"):
            # Update T2 from logs
            t2_acc = log_stream.getvalue()
            
            # Update T1 from report
            for node, state in update.items():
                if "final_report" in state and state["final_report"]:
                    report = state["final_report"]
                    summary = report.get("executive_summary", "")
                    detail = report.get("detailed_breakdown", "")
                    t1_acc = f"EXECUTIVE SUMMARY:\n{summary}\n\nDETAILED REPORT:\n{detail}"
            
            yield t1_acc, t2_acc, gr.update(interactive=False)
            
    except Exception as e:
        yield t1_acc, t2_acc + f"\n[ERROR]: {e}", gr.update(interactive=False)
    finally:
        sys.stdout = old_stdout
        
    yield t1_acc, t2_acc, gr.update(interactive=True)

def handle_approved(content):
    if not content:
        return "No content to save."
    
    try:
        # Construct filename: Title (first line) + Timestamp
        first_line = content.split('\n')[0].strip()[:50]
        safe_title = re.sub(r'[^a-zA-Z0-9]', '_', first_line)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe_title}_{ts}.docx"
        path = REPORTS_DIR / filename
        
        doc = Document()
        doc.add_heading(first_line, 0)
        doc.add_paragraph(content)
        doc.save(path)
        
        # Update Vector DB
        chunk_and_vectorize(REPORTS_DIR)
        
        return f"Saved to reports/{filename} and database updated."
    except Exception as e:
        return f"Error: {e}"

def handle_rag(query):
    if not query: return ""
    yield "Searching knowledge base..."
    res = query_vector_db(query)
    yield str(res)

def toggle_theme():
    # This will be handled by client-side JS for immediate effect
    pass

# ── Interface ─────────────────────────────────────────────────────────────

with gr.Blocks(title="Smart Medical Research Assistant", css=CUSTOM_CSS) as demo:
    # Header
    with gr.Row(elem_id="header-row"):
        gr.Markdown("# SMART MEDICAL RESEARCH ASSISTANT", elem_id="title-markdown")
        theme_btn = gr.Button("🌙", elem_id="theme-toggle", elem_classes="theme-toggle-btn")

    # Main Panels (75/25)
    with gr.Row(elem_classes="main-content-row"):
        with gr.Column(scale=3):
            t1 = gr.Textbox(label="Agent Response", interactive=False, elem_id="t1", lines=15)
            t2 = gr.Textbox(label="Agent Execution Log", interactive=False, elem_id="t2", lines=5)
        with gr.Column(scale=1):
            t4 = gr.Textbox(label="Knowledge Base Response", interactive=False, elem_id="t4", lines=25)

    # Inputs Row
    with gr.Row(elem_classes="inputs-row"):
        with gr.Column(scale=3):
            t3 = gr.Textbox(placeholder="Type something...", show_label=True, label="Your Query", elem_id="t3")
        with gr.Column(scale=1):
            t5 = gr.Textbox(placeholder="Ask something...", show_label=True, label="Knowledge Base Query", elem_id="t5")

    # Buttons Row
    with gr.Row(elem_classes="buttons-row"):
        with gr.Column(scale=3):
            with gr.Row():
                submit_btn = gr.Button("Submit", elem_classes="primary-btn")
                approved_btn = gr.Button("Approve", interactive=False, elem_classes="secondary-btn")
            status_out = gr.Markdown("")
        with gr.Column(scale=1):
            arrow_btn = gr.Button("↑", elem_classes="primary-btn arrow-btn")

    # Client-side Theme Toggle
    theme_btn.click(None, None, None, js="""
        () => {
            document.body.classList.toggle('dark');
            const btn = document.getElementById('theme-toggle');
            btn.innerText = document.body.classList.contains('dark') ? '☀️' : '🌙';
        }
    """)

    # Events
    submit_btn.click(run_pipeline, inputs=t3, outputs=[t1, t2, approved_btn])
    t3.submit(run_pipeline, inputs=t3, outputs=[t1, t2, approved_btn])
    
    approved_btn.click(handle_approved, inputs=t1, outputs=status_out)
    
    arrow_btn.click(handle_rag, inputs=t5, outputs=t4)
    t5.submit(handle_rag, inputs=t5, outputs=t4)

if __name__ == "__main__":
    demo.launch(server_name="127.0.0.1", server_port=7860)
